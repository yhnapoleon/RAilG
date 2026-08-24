"""Office 文档提取器:DOCX / XLSX。

按 extractors/base.py 的 markdown 契约实现:段落样式映射成标题层级,
表格转成 markdown 表格。
"""

from __future__ import annotations

from pathlib import Path

from railg.ingest.extractors.base import ExtractedDocument, Extractor

_MAX_CELL = 2000


def _esc(text: str) -> str:
    return (text or "").replace("|", "\\|").replace("\n", " ").strip()[:_MAX_CELL]


def _rows_to_markdown(rows: list[list[str]]) -> str:
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    head, *body = rows
    out = ["| " + " | ".join(head) + " |", "| " + " | ".join(["---"] * width) + " |"]
    out += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(out)


class DocxExtractor(Extractor):
    name = "docx"
    extensions = (".docx",)

    def extract(self, path: Path) -> ExtractedDocument:
        import docx

        doc = docx.Document(str(path))
        parts: list[str] = []

        # 按文档流顺序遍历,保证表格落在正确的段落之间
        body = doc.element.body
        tables = iter(doc.tables)
        paragraphs = {p._element: p for p in doc.paragraphs}

        for child in body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                para = paragraphs.get(child)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                parts.append(self._style_to_markdown(para.style.name or "", text))
            elif tag == "tbl":
                table = next(tables, None)
                if table is None:
                    continue
                rows = [[_esc(c.text) for c in row.cells] for row in table.rows]
                md = _rows_to_markdown(rows)
                if md:
                    parts.append(md)

        core = doc.core_properties
        meta = {k: v for k, v in {
            "title": core.title, "author": core.author, "subject": core.subject,
        }.items() if v}

        return ExtractedDocument(file_markdown="\n\n".join(parts), meta=meta)

    @staticmethod
    def _style_to_markdown(style: str, text: str) -> str:
        style_low = style.lower()
        if "heading" in style_low or style_low.startswith("标题"):
            digits = "".join(ch for ch in style if ch.isdigit())
            level = int(digits) if digits else 1
            return f"{'#' * min(level, 3)} {text}"
        if style_low.startswith("title"):
            return f"# {text}"
        if "list" in style_low:
            return f"- {text}"
        return text


class XlsxExtractor(Extractor):
    name = "xlsx"
    extensions = (".xlsx", ".xlsm")

    def extract(self, path: Path) -> ExtractedDocument:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        # 每个 sheet 作为一"页",既符合直觉也让 chunker 能给出 sheet 级定位
        pages: list[str] = []
        titles: list[str] = []
        try:
            for ws in wb.worksheets:
                titles.append(ws.title)
                rows = [
                    [_esc("" if c is None else str(c)) for c in row]
                    for row in ws.iter_rows(values_only=True)
                ]
                md = _rows_to_markdown(rows)
                pages.append(f"## {ws.title}\n\n{md}" if md else f"## {ws.title}")
        finally:
            # read_only 模式下关闭后不能再访问 worksheet,元信息必须先收集
            wb.close()

        return ExtractedDocument(
            page_markdowns=pages,
            meta={"sheets": titles} if titles else {},
        )
